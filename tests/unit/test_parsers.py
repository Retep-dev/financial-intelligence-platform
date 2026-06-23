import os
import tempfile

import pytest
from docx import Document
from openpyxl import Workbook

from financial_intelligence_platform.services.ingestion.parsers.parser_router import extract_text
from financial_intelligence_platform.services.ingestion.parsers.txt_parser import extract_txt_text
from financial_intelligence_platform.services.ingestion.parsers.csv_parser import extract_csv_text
from financial_intelligence_platform.services.ingestion.parsers.docx_parser import extract_docx_text
from financial_intelligence_platform.services.ingestion.parsers.xlsx_parser import extract_xlsx_text
from financial_intelligence_platform.services.ingestion.parsers.html_parser import extract_html_text


def test_extract_txt_text():
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8"
    ) as f:
        f.write("Revenue was $1,000,000 in Q1 2024.\n")
        path = f.name

    try:
        text = extract_txt_text(path)
        assert "Revenue" in text
        assert "$" in text
    finally:
        os.unlink(path)


def test_extract_csv_text():
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".csv", delete=False, encoding="utf-8"
    ) as f:
        f.write("date,description,amount\n")
        f.write("2024-01-15,Sales,1000\n")
        path = f.name

    try:
        text = extract_csv_text(path)
        assert "Sales" in text
        assert "1000" in text
    finally:
        os.unlink(path)


def test_extract_html_text():
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".html", delete=False, encoding="utf-8"
    ) as f:
        f.write("<html><body><h1>Report</h1><p>Net income: $500</p></body></html>")
        path = f.name

    try:
        text = extract_html_text(path)
        assert "Report" in text
        assert "Net income" in text
    finally:
        os.unlink(path)


def test_extract_docx_text():
    with tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False
    ) as f:
        path = f.name

    try:
        doc = Document()
        doc.add_heading("Annual Report", level=1)
        doc.add_paragraph("Revenue increased to $1.2B in 2024.")
        doc.save(path)

        text = extract_docx_text(path)
        assert "Annual Report" in text
        assert "Revenue" in text
    finally:
        os.unlink(path)


def test_extract_xlsx_text():
    with tempfile.NamedTemporaryFile(
        suffix=".xlsx", delete=False
    ) as f:
        path = f.name

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Financials"
        ws.append(["Metric", "Value"])
        ws.append(["Revenue", "1000000"])
        wb.save(path)

        text = extract_xlsx_text(path)
        assert "Financials" in text
        assert "Revenue" in text
        assert "1000000" in text
    finally:
        os.unlink(path)


def test_parser_router_txt():
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8"
    ) as f:
        f.write("The company reported (50.5) as a loss.")
        path = f.name

    try:
        text = extract_text(path)
        assert "company" in text
        assert "-50.5" in text
    finally:
        os.unlink(path)


def test_parser_router_unsupported():
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".unsupported", delete=False
    ) as f:
        path = f.name

    try:
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(path)
    finally:
        os.unlink(path)
