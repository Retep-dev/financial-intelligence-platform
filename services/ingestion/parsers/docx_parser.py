from docx import Document


def extract_docx_text(file_path: str) -> str:
    doc = Document(file_path)

    paragraphs = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs)
